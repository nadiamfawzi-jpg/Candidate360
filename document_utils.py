from docx import Document
from pypdf import PdfReader


def clean_pdf_text(text):
    """Remove known standalone icon-font artifacts from extracted CV text."""
    cleaned_lines = []

    for line in text.splitlines():
        cleaned_line = " ".join(line.split()).strip()

        # Some CV templates use icon fonts for contact details. PDF text
        # extraction can turn those icons into isolated N or F characters.
        if cleaned_line in {"N", "F"}:
            continue

        if cleaned_line:
            cleaned_lines.append(cleaned_line)

    return "\n".join(cleaned_lines)


def read_pdf(uploaded_file):
    reader = PdfReader(uploaded_file)
    text = ""

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"

    return clean_pdf_text(text)


def read_docx(uploaded_file):
    document = Document(uploaded_file)
    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    # Some CV templates put the skills or contact-details section in a
    # table rather than plain paragraphs. Without this, that content is
    # silently missing from the CV Analyst's input.
    for table in document.tables:
        for row in table.rows:
            row_cells = []
            previous_cell_text = None
            for cell in row.cells:
                cell_text = cell.text.strip()
                # Merged cells repeat the same text for every column they
                # span, so skip immediate repeats rather than doubling them.
                if cell_text and cell_text != previous_cell_text:
                    row_cells.append(cell_text)
                previous_cell_text = cell_text
            if row_cells:
                text += " | ".join(row_cells) + "\n"

    return text


def read_cv(uploaded_file):
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()
    if file_name.endswith(".pdf"):
        return read_pdf(uploaded_file)
    if file_name.endswith(".docx"):
        return read_docx(uploaded_file)
    if file_name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

    raise ValueError("Please upload a PDF, DOCX or TXT file.")
